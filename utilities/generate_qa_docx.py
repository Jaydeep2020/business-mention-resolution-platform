import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from datetime import datetime

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_qa_item(doc, q_num, question, answer_paragraphs, highlights=None):
    # Question Heading
    p_q = doc.add_paragraph()
    p_q.paragraph_format.space_before = Pt(12)
    p_q.paragraph_format.space_after = Pt(4)
    p_q.paragraph_format.keep_with_next = True
    
    r_tag = p_q.add_run(f"Q{q_num}: ")
    r_tag.font.name = 'Calibri'
    r_tag.font.size = Pt(11.5)
    r_tag.font.bold = True
    r_tag.font.color.rgb = RGBColor(30, 58, 138)
    
    r_q = p_q.add_run(question)
    r_q.font.name = 'Calibri'
    r_q.font.size = Pt(11.5)
    r_q.font.bold = True
    r_q.font.color.rgb = RGBColor(15, 23, 42)
    
    # Answer Content
    for p_text in answer_paragraphs:
        p_ans = doc.add_paragraph()
        p_ans.paragraph_format.space_before = Pt(2)
        p_ans.paragraph_format.space_after = Pt(4)
        p_ans.paragraph_format.left_indent = Inches(0.2)
        r_ans = p_ans.add_run(p_text)
        r_ans.font.name = 'Calibri'
        r_ans.font.size = Pt(10)
        r_ans.font.color.rgb = RGBColor(30, 41, 59)
        
    if highlights:
        # Callout box for highlights
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, "F0F9FF")
        set_cell_margins(cell, top=80, bottom=80, left=150, right=150)
        
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:left w:val="single" w:sz="18" w:space="0" w:color="0284C7"/>'
            f'<w:top w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'<w:bottom w:val="none"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)
        
        p_hl = cell.paragraphs[0]
        p_hl.paragraph_format.space_before = Pt(0)
        p_hl.paragraph_format.space_after = Pt(2)
        r_hl_t = p_hl.add_run("💡 Key Takeaway: ")
        r_hl_t.font.bold = True
        r_hl_t.font.color.rgb = RGBColor(3, 105, 161)
        r_hl_text = p_hl.add_run(highlights)
        r_hl_text.font.size = Pt(9.5)
        r_hl_text.font.color.rgb = RGBColor(30, 41, 59)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def generate_qa_docx():
    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)
        
    # Title Block
    p_t = doc.add_paragraph()
    p_t.paragraph_format.space_before = Pt(0)
    p_t.paragraph_format.space_after = Pt(2)
    r_t = p_t.add_run("Project Evaluation & Technical Interview Defense Guide")
    r_t.font.name = 'Calibri'
    r_t.font.size = Pt(20)
    r_t.font.bold = True
    r_t.font.color.rgb = RGBColor(30, 58, 138)
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(12)
    r_sub = p_sub.add_run("Business Mention Resolution Platform (BMRP) — Architecture, Algorithms & What-If Scenarios")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(100, 116, 139)
    
    # ---------------------------------------------------------
    # SECTION 1
    # ---------------------------------------------------------
    p_sec1 = doc.add_heading("1. High-Level Project Overview & Business Value", level=1)
    p_sec1.paragraph_format.space_before = Pt(14)
    p_sec1.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    add_qa_item(
        doc, "1.1",
        "What did you build? Give a 2-minute elevator pitch.",
        [
            "The Business Mention Resolution Platform (BMRP) is an enterprise AI microservice system that extracts, disambiguates, and resolves local business mentions in unstructured user reviews to canonical records in a catalog of ~150,000 businesses (Yelp Academic Dataset).",
            "The system features a multi-stage hybrid pipeline: zero-shot NER with GLiNER, hybrid SQL + FAISS vector candidate search (98.8% Recall@20), multi-factor scoring (50% name, 25% embeddings, 25% geography), a LangGraph Agentic Assistant for complex edge-cases, and ReportLab PDF document generation with a Streamlit human review portal."
        ],
        "Emphasize the hybrid design: high-speed deterministic rules for clear matches + selective LLM reasoning for edge cases."
    )
    
    add_qa_item(
        doc, "1.2",
        "Why is local business entity resolution hard?",
        [
            "1. Polysemy & Chain Stores: Chains like Starbucks or Domino's share identical names across thousands of locations. Resolving requires pinpointing the exact street and city.",
            "2. Noisy Text: Users omit legal suffixes ('LLC', 'Inc.'), use abbreviations, or make typographical errors.",
            "3. High Automation Risk: Incorrectly linking negative reviews corrupts CRM intelligence and sentiment analysis."
        ]
    )
    
    # ---------------------------------------------------------
    # SECTION 2
    # ---------------------------------------------------------
    p_sec2 = doc.add_heading("2. NLP & Mention Extraction (GLiNER vs. spaCy)", level=1)
    p_sec2.paragraph_format.space_before = Pt(14)
    p_sec2.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    add_qa_item(
        doc, "2.1",
        "Why did spaCy fail and why did you choose GLiNER?",
        [
            "spaCy's general-purpose en_core_web_sm model frequently misclassifies single brand names like 'Domino' or 'Tony' as PERSON. If we filter only for ORG, brand mentions are missed. If we allow PERSON, normal human names become false business mentions.",
            "GLiNER is a zero-shot bidirectional transformer where we explicitly define target labels: ['business', 'restaurant', 'store', 'hotel', 'cafe']. It improved our extraction F1-score from 64.1% (spaCy) to 91.1% without requiring labeled training datasets."
        ],
        "GLiNER also uses negative context labels ('city', 'state', 'address') during inference to suppress geographic false positives."
    )
    
    # ---------------------------------------------------------
    # SECTION 3
    # ---------------------------------------------------------
    p_sec3 = doc.add_heading("3. Candidate Generation & Scoring Math", level=1)
    p_sec3.paragraph_format.space_before = Pt(14)
    p_sec3.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    add_qa_item(
        doc, "3.1",
        "Explain the candidate scoring formula and weights.",
        [
            "Formula: Final Score = 0.50 * S_name + 0.25 * S_embed + 0.10 * S_city + 0.05 * S_state + 0.10 * S_addr",
            "• S_name (50%): 0.70 * SequenceMatcher + 0.30 * Jaccard Token Overlap.",
            "• S_embed (25%): Normalized cosine similarity of 384-d sentence-transformers/all-MiniLM-L6-v2 embeddings via FAISS.",
            "• Geography (25%): City (10%), State (5%), and Street Address (10%) contextual extraction."
        ],
        "Lexical brand identity is the strongest signal (50%), preventing competitor false matches even when categories are identical."
    )
    
    # ---------------------------------------------------------
    # SECTION 4
    # ---------------------------------------------------------
    p_sec4 = doc.add_heading("4. LangGraph AI Assistant & Decisioning", level=1)
    p_sec4.paragraph_format.space_before = Pt(14)
    p_sec4.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    add_qa_item(
        doc, "4.1",
        "What are your auto-resolution rules and how does LangGraph work?",
        [
            "Auto-Resolution Criteria: Score >= 0.85, Ambiguity Gap >= 0.05, and is_verified == True.",
            "LangGraph Workflow: load_mention -> generate_candidates -> assess_candidates -> (forced_review / direct_resolve / analyze_context) -> validate_recommendation -> persist_decision.",
            "Guardrail Policy: Unverified businesses deterministically route to human review (never auto-resolve). Obvious verified matches resolve in <10ms with 0 token cost."
        ],
        "Pydantic structured output validation guarantees the LLM cannot hallucinate invalid database IDs."
    )
    
    # ---------------------------------------------------------
    # SECTION 5
    # ---------------------------------------------------------
    p_sec5 = doc.add_heading("5. Crucial 'What If' & Scalability Scenarios", level=1)
    p_sec5.paragraph_format.space_before = Pt(14)
    p_sec5.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    add_qa_item(
        doc, "5.1",
        "What if the catalog scales from 150K to 50 Million businesses?",
        [
            "1. FAISS: Transition from IndexFlatIP to IndexIVFFlat / IndexHNSW with Product Quantization (PQ8) for <2ms search across 50M records.",
            "2. Lexical: Replace PostgreSQL ILIKE with Elasticsearch / OpenSearch BM25 cluster.",
            "3. Database: Range-partition PostgreSQL tables by country/state."
        ]
    )
    
    add_qa_item(
        doc, "5.2",
        "What if we ingest a real-time stream of 10,000 reviews/second?",
        [
            "1. Buffer stream via Apache Kafka / RabbitMQ topics.",
            "2. Deploy asynchronous Celery/Ray worker pools with batch GPU inference on GLiNER and MiniLM embeddings (batch size 128).",
            "3. Bulk write resolution results using PostgreSQL COPY protocol."
        ]
    )
    
    add_qa_item(
        doc, "5.3",
        "What if OpenAI API is down or rate-limited?",
        [
            "68% of mentions resolve deterministically without OpenAI. For edge cases, self-host an open-weights SLM (Mistral-7B / Llama-3-8B) via vLLM locally by reconfiguring OPENAI_BASE_URL."
        ]
    )
    
    out_path = r"D:\Sculptsoft\business-mention-resolution-platform\PROJECT_EVALUATION_QA.docx"
    doc.save(out_path)
    print(f"Generated Q&A Word doc at: {out_path}")

if __name__ == "__main__":
    generate_qa_docx()
